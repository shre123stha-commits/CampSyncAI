from pathlib import Path
from docx import Document


def read_docx(file_path: Path) -> str:
    """
    Reads all text from a Word document.
    """

    doc = Document(file_path)

    text = []

    # Read paragraphs
    for para in doc.paragraphs:
        if para.text.strip():
            text.append(para.text.strip())

    # Read tables
    for table in doc.tables:
        for row in table.rows:
            row_text = [
                cell.text.strip()
                for cell in row.cells
            ]
            text.append(" | ".join(row_text))

    return "\n".join(text)